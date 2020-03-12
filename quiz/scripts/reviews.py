from django.core.mail import EmailMessage
from quiz.models import MyUser, UserAnswer, Answer
from docx import Document
from docx.shared import Inches

def run():
	'''
	This script creates an analysis of the answers given by the participants and mails it to them.
	The analysis contains question, options, given answer and actual answer.
	'''
	userAnswers = UserAnswer.objects.all()
	try:
		for userAnswer in userAnswers:
			document = Document()
			document.add_heading('MCQ Review', 0)
			questions = userAnswer.questions.split(',')
			answers = userAnswer.answers.split(',')
			for i,question_id in enumerate(questions):
				question = Answer.objects.get(question_id=question_id)
				document.add_heading('Question '+str(i+1)+'. '+str(question.question), level=3)
				document.add_paragraph('')
				document.add_paragraph('A. '+question.option_a)
				document.add_paragraph('B. '+question.option_b)
				document.add_paragraph('C. '+question.option_c)
				document.add_paragraph('D. '+question.option_d)
				document.add_paragraph('Correct answers : '+question.correct_answer)
				document.add_paragraph('Your answer : '+answers[i])
				document.add_page_break()
			document.save('review.docx')
			email_id = []
			if userAnswer.user.email_2:
				email_id = [userAnswer.user.email_1, userAnswer.user.email_2]
			else:
				email_id = [userAnswer.user.email_1]
			email = EmailMessage(
				'Review for Quiz',
				'PFA file containing review for MCQ Quiz', 
				'Team Django-Quiz',
				email_id,
			)
			email.attach_file('review.docx')
			email.send(fail_silently=False)
	except:
		pass
